"""
Check the built report against the format rules in the assessment brief:

    Paper A4 | Margins 1.5" left, 1" right/top/bottom
    Page numbers - bottom, right | Line spacing 1.5
    Headings 14pt bold | Normal 12pt | Times New Roman

Reads the .docx for the settings and the .pdf for what actually came out.
"""
import os, sys
from docx import Document
from docx.shared import Pt
import fitz

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(HERE, "CIS6003_WRIT1_Report.docx")
PDF = os.path.join(HERE, "CIS6003_WRIT1_Report.pdf")

PASS, FAIL = "PASS", "FAIL"
results = []


def check(label, ok, detail=""):
    results.append((PASS if ok else FAIL, label, detail))


def resolved_font(style):
    """
    The font a style actually renders in, following the base-style chain.

    Word drops a style's own rFonts when it is identical to the one it
    inherits, so reading style.font.name alone reports None for a heading
    that is correctly set to Times New Roman via Normal.
    """
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        if style.font.name:
            return style.font.name
        style = style.base_style
    return None


def main():
    doc = Document(DOCX)

    # ---------------- page setup ----------------
    bad_size = bad_margin = 0
    for s in doc.sections:
        w, h = sorted((s.page_width.inches, s.page_height.inches))
        if not (abs(w - 8.268) < 0.02 and abs(h - 11.693) < 0.02):
            bad_size += 1
        if not (abs(s.left_margin.inches - 1.5) < 0.01
                and abs(s.right_margin.inches - 1.0) < 0.01
                and abs(s.top_margin.inches - 1.0) < 0.01
                and abs(s.bottom_margin.inches - 1.0) < 0.01):
            bad_margin += 1
    check("A4 paper in every section", bad_size == 0,
          f"{len(doc.sections)} sections, {bad_size} wrong")
    check("Margins 1.5\" left / 1\" right, top, bottom", bad_margin == 0,
          f"{bad_margin} sections wrong")

    # ---------------- fonts ----------------
    normal = doc.styles["Normal"]
    check("Normal font is Times New Roman", normal.font.name == "Times New Roman",
          str(normal.font.name))
    check("Normal size is 12pt", normal.font.size == Pt(12),
          f"{normal.font.size.pt}pt")
    check("Normal line spacing is 1.5",
          abs(normal.paragraph_format.line_spacing - 1.5) < 0.001,
          str(normal.paragraph_format.line_spacing))

    for lvl in (1, 2, 3):
        st = doc.styles[f"Heading {lvl}"]
        check(f"Heading {lvl}: {resolved_font(st)}, 14pt, bold",
              resolved_font(st) == "Times New Roman"
              and st.font.size == Pt(14) and st.font.bold,
              f"{st.font.size.pt}pt, bold={st.font.bold}")

    # ---------------- footers ----------------
    with_field = 0
    for s in doc.sections:
        xml = s.footer.paragraphs[0]._p.xml if s.footer.paragraphs else ""
        if "PAGE" in xml:
            with_field += 1
    check("Page-number field in every section footer",
          with_field == len(doc.sections), f"{with_field}/{len(doc.sections)}")

    # ---------------- the PDF that came out ----------------
    if os.path.exists(PDF):
        d = fitz.open(PDF)
        blanks = [i + 1 for i, p in enumerate(d)
                  if len(p.get_text().strip()) < 6 and not p.get_images()]
        check("No blank pages in the PDF", not blanks, str(blanks))

        missing = []
        for i, p in enumerate(d):
            r = p.rect
            foot = [w for w in p.get_text("words")
                    if w[3] > r.height - 72 and w[2] > r.width * 0.65]
            if not any(w[4].strip().isdigit() for w in foot):
                missing.append(i + 1)
        check("Page number bottom-right on every page", not missing,
              f"missing on {missing}" if missing else f"{d.page_count} pages")

        land = [i + 1 for i, p in enumerate(d) if p.rect.width > p.rect.height]
        check("Landscape pages used only for wide diagrams", len(land) == 4,
              f"pages {land}")

        # What the reader actually sees: the fonts embedded in the PDF.
        fonts = {}
        for p in d:
            for b in p.get_text("dict")["blocks"]:
                for l in b.get("lines", []):
                    for s in l["spans"]:
                        if s["text"].strip():
                            fonts[s["font"]] = fonts.get(s["font"], 0) + 1
        # SymbolMT is Word's glyph for the list bullet, not body text.
        allowed = ("Times", "Consolas", "SymbolMT")
        stray = {f: n for f, n in fonts.items()
                 if not any(a in f for a in allowed)}
        check("PDF text is Times New Roman (Consolas for code, "
              "SymbolMT for bullets)",
              not stray, str(stray) if stray else ", ".join(sorted(fonts)))

        head = [s["size"] for p in d for b in p.get_text("dict")["blocks"]
                for l in b.get("lines", []) for s in l["spans"]
                if "Bold" in s["font"] and round(s["size"]) == 14]
        check("14pt bold headings present in the PDF", len(head) > 20,
              f"{len(head)} runs at 14pt bold")
        print(f"\nPDF: {d.page_count} pages\n")
    else:
        check("PDF exists", False, "run export_pdf.py")

    # ---------------- report ----------------
    width = max(len(l) for _, l, _ in results) + 2
    for status, label, detail in results:
        print(f"[{status}] {label:<{width}} {detail}")
    failed = sum(1 for s, _, _ in results if s == FAIL)
    print(f"\n{len(results) - failed}/{len(results)} checks passed")
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
