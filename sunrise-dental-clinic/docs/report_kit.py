"""
Word-document helpers that enforce the format the brief asks for:

    Paper A4 | Margins 1.5" left, 1" right/top/bottom
    Page numbers - bottom, right | Line spacing 1.5
    Headings 14pt bold | Normal 12pt | Times New Roman

Every helper below writes through these settings rather than around them, so
the finished document cannot drift away from the specification.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FONT = "Times New Roman"
BODY_PT = 12
HEAD_PT = 14
A4_W, A4_H = Cm(21.0), Cm(29.7)

# usable text width in a portrait section with 1.5" + 1" margins
TEXT_W_PORTRAIT = Inches(8.268 - 1.5 - 1.0)      # 5.768"
TEXT_W_LANDSCAPE = Inches(11.693 - 1.5 - 1.0)    # 9.193"


# --------------------------------------------------------------------- #
#  low-level XML helpers
# --------------------------------------------------------------------- #
def _field(paragraph, instr):
    """Insert a Word field (PAGE, TOC, SEQ ...) that Word can refresh."""
    run = paragraph.add_run()
    for kind, text in (("begin", None), ("instrText", instr), ("end", None)):
        el = OxmlElement(f"w:fld{'Char' if kind != 'instrText' else ''}"
                         if kind != "instrText" else "w:instrText")
        if kind == "instrText":
            el.set(qn("xml:space"), "preserve")
            el.text = text
        else:
            el.set(qn("w:fldCharType"), kind)
        run._r.append(el)
    return run


def set_cell_bg(cell, hex_colour):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def force_font(style, name=None):
    """
    Pin a style to a real font face.

    Word's built-in heading styles reference the document theme
    (w:asciiTheme="majorHAnsi"), and a theme reference wins over w:ascii - so
    assigning style.font.name alone leaves the style still following the
    theme. The theme attributes have to be removed as well.
    """
    name = name or FONT
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        if rfonts.get(qn(f"w:{attr}")) is not None:
            del rfonts.attrib[qn(f"w:{attr}")]
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    return style


# --------------------------------------------------------------------- #
#  document set-up
# --------------------------------------------------------------------- #
def new_document():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    force_font(normal)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    for name, size, bold, before, after in [
            ("Heading 1", HEAD_PT, True, 18, 10),
            ("Heading 2", HEAD_PT, True, 14, 8),
            ("Heading 3", HEAD_PT, True, 12, 6),
            ("Heading 4", BODY_PT, True, 10, 4)]:
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.italic = False
        st.font.color.rgb = RGBColor(0, 0, 0)
        force_font(st)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.keep_with_next = True

    cap = doc.styles["Caption"]
    cap.font.name = FONT
    cap.font.size = Pt(10)
    cap.font.bold = False
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    force_font(cap)

    # The TOC styles Word generates also default to the theme font.
    for i in range(1, 4):
        try:
            force_font(doc.styles[f"TOC {i}"])
        except KeyError:
            pass

    portrait(doc.sections[0])
    return doc


def portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = A4_W, A4_H
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    return section


def landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = A4_H, A4_W
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    return section


def page_numbers(doc, start_section=0):
    """Bottom-right page number in every section from start_section on."""
    for i, section in enumerate(doc.sections):
        footer = section.footer
        if i > 0:
            footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        if i >= start_section:
            _field(p, "PAGE  \\* MERGEFORMAT")
        for r in p.runs:
            r.font.name = FONT
            r.font.size = Pt(BODY_PT)


# --------------------------------------------------------------------- #
#  content helpers
# --------------------------------------------------------------------- #
def para(doc, text="", size=BODY_PT, bold=False, italic=False,
         align="justify", space_after=6, space_before=0, indent=None,
         color=None, spacing=1.5, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = spacing
    if indent is not None:
        pf.left_indent = Inches(indent)
    if text:
        add_rich(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_rich(p, text, size=BODY_PT, bold=False, italic=False, color=None):
    """Supports **bold** and `mono` inline markers, nothing else."""
    import re
    for chunk in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not chunk:
            continue
        r = p.add_run()
        if chunk.startswith("**") and chunk.endswith("**"):
            r.text = chunk[2:-2]
            r.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r.text = chunk[1:-1]
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            r.font.size = Pt(size - 1.5)
            r.bold = bold
            continue
        else:
            r.text = chunk
            r.bold = bold
        r.italic = italic
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def bullet(doc, text, level=0, size=BODY_PT):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.3 * level)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    add_rich(p, text, size=size)
    return p


def numbered(doc, text, level=0, size=BODY_PT):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.3 + 0.3 * level)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    add_rich(p, text, size=size)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.size = Pt(HEAD_PT if level <= 3 else BODY_PT)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def question_box(doc, title, lines):
    """The task wording, reproduced from the assessment brief."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, "F2F5F7")
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    add_rich(p, title, bold=True)
    for ln in lines:
        q = cell.add_paragraph()
        q.paragraph_format.line_spacing = 1.5
        q.paragraph_format.space_after = Pt(4)
        q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_rich(q, ln, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def table(doc, headers, rows, widths=None, font_size=10.5, header_bg="D8ECEE",
          landscape_mode=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    total = TEXT_W_LANDSCAPE if landscape_mode else TEXT_W_PORTRAIT
    if widths:
        s = sum(widths)
        widths = [Emu(int(total * w / s)) for w in widths]
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, header_bg)
        p = c.paragraphs[0]
        p.text = ""
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        add_rich(p, h, size=font_size, bold=True)
        if widths:
            c.width = widths[i]
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            p = cells[i].paragraphs[0]
            p.text = ""
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(2)
            add_rich(p, str(v), size=font_size)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(doc, path, caption, width=None, landscape_mode=False):
    """A centred picture with an italic, numbered caption underneath."""
    if width is None:
        width = TEXT_W_LANDSCAPE if landscape_mode else TEXT_W_PORTRAIT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(path, width=width)
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    c.paragraph_format.line_spacing = 1.0
    add_rich(c, caption, size=10, italic=True)
    for r in c.runs:
        r.font.name = FONT
        r.italic = True
        r.font.size = Pt(10)
    return p


def toc_field(doc, levels="1-3"):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    _field(p, f'TOC \\o "{levels}" \\h \\z \\u')
    return p


def full_page_figure(doc, path, caption, orient="landscape"):
    """
    Put one figure on a page of its own, in its own section, so a wide diagram
    gets the whole sheet and is actually readable.

    A landscape A4 page only has 6.27" of usable height, so a 4:3-ish diagram
    scaled to the full 9.19" text width would be 6.5" tall and spill onto a
    second, blank page. The width is therefore derived from the height the
    page can actually give, not from the width.
    """
    from PIL import Image

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    is_land = orient == "landscape"
    (landscape if is_land else portrait)(sec)

    with Image.open(path) as im:
        aspect = im.width / im.height

    if is_land:
        max_h = 8.268 - 2.0 - 0.55          # page height - margins - caption
        max_w = 11.693 - 1.5 - 1.0
    else:
        max_h = 11.693 - 2.0 - 0.55
        max_w = 8.268 - 1.5 - 1.0

    width = min(max_w, max_h * aspect)
    figure(doc, path, caption, width=Inches(width), landscape_mode=is_land)
    return sec


def resume_body(doc):
    """
    Return to portrait after a run of full-page figures.

    Call this once at the end of a run rather than after every figure: each
    section break starts a new page, so restoring portrait between two
    consecutive figures leaves an empty page behind.
    """
    return portrait(doc.add_section(WD_SECTION.NEW_PAGE))
